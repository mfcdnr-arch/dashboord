"""Форматные парсеры документов.

Каждый парсер принимает байты файла и возвращает `ParseResult` — список
таблиц (сетка строк, ячейки как строки) + диапазоны объединённых ячеек
+ предупреждения. Никакого ввода-вывода и обращений к БД: чистая обработка
байтов, легко тестировать.

Форматы v1 (док-06): Excel (.xlsx/.xls), CSV, Word (.docx), PDF (текстовый слой).
OCR не поддерживается: для PDF без текста возвращаем предупреждение.

Геометрия сетки намеренно повторяет оригинал: пустые строки ВНУТРИ и НАД
таблицей сохраняются, обрезаются только хвостовые. Иначе номера строк
разойдутся с файлом, и предпросмотр «как в оригинале» построить нельзя.
Значение объединённой ячейки лежит ТОЛЬКО в её левой верхней клетке, а сам
диапазон — в `merges`: так сетку можно нарисовать через rowspan/colspan,
а для анализа значений развернуть через `fill_merges`.
"""
from __future__ import annotations

import csv
import datetime as dt
import io
from dataclasses import dataclass, field
from typing import List, Optional, Sequence, Tuple

# Диапазон объединения: (первая строка, первый столбец, последняя строка,
# последний столбец) — включительно, нумерация с нуля.
Merge = Tuple[int, int, int, int]


class UnsupportedFormat(Exception):
    """Формат файла не поддерживается парсером."""


@dataclass
class ParsedTable:
    """Одна распознанная табличная область (лист/страница)."""

    sheet_or_page: Optional[str]
    table_index: int
    rows: List[List[str]]  # все строки, включая шапку; ячейки — строки ("" если пусто)
    merges: List[Merge] = field(default_factory=list)

    @property
    def row_count(self) -> int:
        return len(self.rows)

    @property
    def column_count(self) -> int:
        return max((len(r) for r in self.rows), default=0)


@dataclass
class ParseResult:
    tables: List[ParsedTable] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)


# --------------------------------------------------------------------------- #
# Нормализация значений ячеек
# --------------------------------------------------------------------------- #
def _cell_to_str(value) -> str:
    """Приводит значение ячейки к строке (устойчиво к типам Excel/дат)."""
    if value is None:
        return ""
    if isinstance(value, bool):
        return "да" if value else "нет"
    if isinstance(value, (dt.datetime,)):
        # полночь → только дата, иначе дата+время
        if value.hour == 0 and value.minute == 0 and value.second == 0:
            return value.date().isoformat()
        return value.isoformat(sep=" ")
    if isinstance(value, dt.date):
        return value.isoformat()
    if isinstance(value, float):
        # целочисленные float без хвоста ".0"
        if value.is_integer():
            return str(int(value))
        return repr(value)
    return str(value).strip()


def _normalize_grid(
    grid: List[List], merges: Optional[Sequence[Merge]] = None
) -> Tuple[List[List[str]], List[Merge]]:
    """Строки → строковые ячейки; обрезает ТОЛЬКО хвостовые пустые строки/столбцы.

    Пустые строки внутри таблицы и над ней сохраняются: по ним пользователь
    сверяет предпросмотр с оригиналом, а нумерация строк должна совпадать с
    Excel. Хвост обрезаем, иначе фантомные строки листа раздувают сетку.
    Диапазоны объединений подрезаются по новым границам.
    """
    rows = [[_cell_to_str(c) for c in row] for row in grid]
    # выравниваем по максимальной ширине
    width = max((len(r) for r in rows), default=0)
    rows = [r + [""] * (width - len(r)) for r in rows]
    # убрать хвостовые пустые столбцы
    while width > 0 and all(r[width - 1] == "" for r in rows):
        for r in rows:
            r.pop()
        width -= 1
    # убрать хвостовые пустые строки
    while rows and all(c == "" for c in rows[-1]):
        rows.pop()

    height = len(rows)
    clipped: List[Merge] = []
    for r1, c1, r2, c2 in merges or ():
        if r1 >= height or c1 >= width:
            continue  # объединение целиком в отрезанном хвосте
        r2, c2 = min(r2, height - 1), min(c2, width - 1)
        if r2 > r1 or c2 > c1:  # вырожденное 1×1 после подрезки не храним
            clipped.append((r1, c1, r2, c2))
    return rows, clipped


def fill_merges(rows: List[List[str]], merges: Sequence[Merge]) -> List[List[str]]:
    """Копия сетки, где значение объединения размножено на весь его диапазон.

    Нужна для АНАЛИЗА и материализации значений (название строки, объединённое
    на несколько строк, относится к каждой из них). Для отрисовки предпросмотра
    использовать НЕ надо — там объединения рисуются через rowspan/colspan.
    """
    out = [list(r) for r in rows]
    for r1, c1, r2, c2 in merges:
        if r1 >= len(out) or c1 >= len(out[r1]):
            continue
        value = out[r1][c1]
        if value == "":
            continue
        for r in range(r1, min(r2, len(out) - 1) + 1):
            for c in range(c1, min(c2, len(out[r]) - 1) + 1):
                out[r][c] = value
    return out


# --------------------------------------------------------------------------- #
# Excel .xlsx (openpyxl)
# --------------------------------------------------------------------------- #
def parse_xlsx(content: bytes) -> ParseResult:
    from openpyxl import load_workbook

    result = ParseResult()
    wb = load_workbook(io.BytesIO(content), read_only=False, data_only=True)
    try:
        for idx, ws in enumerate(wb.worksheets):
            grid = [list(row) for row in ws.iter_rows(values_only=True)]
            # openpyxl отдаёт значение объединения только в левой верхней клетке —
            # ровно то, что нужно для отрисовки; тиражирование делает fill_merges.
            raw_merges = [
                (rng.min_row - 1, rng.min_col - 1, rng.max_row - 1, rng.max_col - 1)
                for rng in ws.merged_cells.ranges
            ]
            rows, merges = _normalize_grid(grid, raw_merges)
            if rows:
                result.tables.append(
                    ParsedTable(sheet_or_page=ws.title, table_index=idx, rows=rows, merges=merges)
                )
        if not result.tables:
            result.warnings.append("В книге Excel не найдено непустых листов.")
    finally:
        wb.close()
    return result


# --------------------------------------------------------------------------- #
# Excel .xls (xlrd)
# --------------------------------------------------------------------------- #
def parse_xls(content: bytes) -> ParseResult:
    import xlrd

    result = ParseResult()
    # formatting_info нужен ради merged_cells; часть .xls его не отдаёт —
    # тогда работаем без объединений, а не падаем.
    try:
        book = xlrd.open_workbook(file_contents=content, formatting_info=True)
    except NotImplementedError:
        book = xlrd.open_workbook(file_contents=content)
    for idx in range(book.nsheets):
        sheet = book.sheet_by_index(idx)
        grid: List[List] = []
        for r in range(sheet.nrows):
            grid.append([_xls_cell(book, sheet.cell(r, c)) for c in range(sheet.ncols)])
        # xlrd отдаёт (rlo, rhi, clo, chi) с ИСКЛЮЧАЮЩИМИ верхними границами
        raw_merges = [
            (rlo, clo, rhi - 1, chi - 1) for rlo, rhi, clo, chi in getattr(sheet, "merged_cells", [])
        ]
        # в отличие от openpyxl, xlrd держит значение только в левой верхней —
        # дополнительная зачистка не нужна
        rows, merges = _normalize_grid(grid, raw_merges)
        if rows:
            result.tables.append(
                ParsedTable(sheet_or_page=sheet.name, table_index=idx, rows=rows, merges=merges)
            )
    if not result.tables:
        result.warnings.append("В книге Excel (.xls) не найдено непустых листов.")
    return result


def _xls_cell(book, cell):
    import xlrd

    if cell.ctype == xlrd.XL_CELL_DATE:
        try:
            return dt.datetime(*xlrd.xldate_as_tuple(cell.value, book.datemode))
        except Exception:
            return cell.value
    if cell.ctype == xlrd.XL_CELL_BOOLEAN:
        return bool(cell.value)
    return cell.value


# --------------------------------------------------------------------------- #
# CSV (автоопределение кодировки и разделителя)
# --------------------------------------------------------------------------- #
def parse_csv(content: bytes) -> ParseResult:
    from charset_normalizer import from_bytes

    result = ParseResult()
    match = from_bytes(content).best()
    if match is None:
        text = content.decode("utf-8", errors="replace")
        result.warnings.append("Не удалось определить кодировку CSV — прочитано как UTF-8.")
    else:
        text = str(match)
        if match.encoding and match.encoding.lower() not in ("utf_8", "utf-8", "ascii"):
            result.warnings.append(f"Кодировка CSV определена как {match.encoding} — подтвердите.")

    sample = text[:8192]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        delimiter = dialect.delimiter
    except csv.Error:
        delimiter = ";" if sample.count(";") > sample.count(",") else ","
        result.warnings.append(f"Разделитель CSV определён эвристикой как «{delimiter}» — подтвердите.")

    reader = csv.reader(io.StringIO(text), delimiter=delimiter)
    rows, _ = _normalize_grid([list(r) for r in reader])  # в CSV объединений не бывает
    if rows:
        result.tables.append(ParsedTable(sheet_or_page=None, table_index=0, rows=rows))
    else:
        result.warnings.append("CSV не содержит данных.")
    return result


# --------------------------------------------------------------------------- #
# Word .docx (таблицы; текст вне таблиц игнорируется в v1)
# --------------------------------------------------------------------------- #
def parse_docx(content: bytes) -> ParseResult:
    from docx import Document

    result = ParseResult()
    doc = Document(io.BytesIO(content))
    for idx, table in enumerate(doc.tables):
        grid, raw_merges = _docx_grid(table)
        rows, merges = _normalize_grid(grid, raw_merges)
        if rows:
            result.tables.append(
                ParsedTable(
                    sheet_or_page=f"Таблица {idx + 1}", table_index=idx, rows=rows, merges=merges
                )
            )
    if not result.tables:
        result.warnings.append("В документе Word не найдено таблиц (текст вне таблиц в v1 не извлекается).")
    return result


def _docx_grid(table) -> Tuple[List[List[str]], List[Merge]]:
    """Сетка таблицы Word + диапазоны объединений.

    python-docx возвращает объединённую ячейку столько раз, сколько клеток она
    занимает (и текст повторяется в каждой). Опознаём объединения по одному и
    тому же XML-элементу `_tc`, текст оставляем только в левой верхней клетке —
    как в остальных форматах. Ключом словаря служит сам элемент lxml (сравнение
    по тождеству), `id()` для этого не годится: прокси-объекты пересоздаются.
    """
    grid: List[List[str]] = []
    box: dict = {}
    for r, row in enumerate(table.rows):
        cells = list(row.cells)
        grid.append([c.text for c in cells])
        for c, cell in enumerate(cells):
            b = box.get(cell._tc)
            if b is None:
                box[cell._tc] = [r, c, r, c]
            else:
                b[0], b[1] = min(b[0], r), min(b[1], c)
                b[2], b[3] = max(b[2], r), max(b[3], c)

    merges: List[Merge] = []
    for r1, c1, r2, c2 in box.values():
        if r2 == r1 and c2 == c1:
            continue
        merges.append((r1, c1, r2, c2))
        for r in range(r1, r2 + 1):
            for c in range(c1, c2 + 1):
                if (r, c) != (r1, c1) and r < len(grid) and c < len(grid[r]):
                    grid[r][c] = ""
    return grid, merges


# --------------------------------------------------------------------------- #
# PDF (текстовый слой; OCR вне MVP)
# --------------------------------------------------------------------------- #
def parse_pdf(content: bytes) -> ParseResult:
    import pdfplumber

    result = ParseResult()
    has_text = False
    table_index = 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            if (page.extract_text() or "").strip():
                has_text = True
            for tbl in page.extract_tables():
                # pdfplumber не сообщает об объединениях — они восстановимы
                # только по геометрии линий, это вне текущего объёма
                rows, _ = _normalize_grid([list(r) for r in tbl])
                if rows:
                    result.tables.append(
                        ParsedTable(sheet_or_page=f"Стр. {page_no}", table_index=table_index, rows=rows)
                    )
                    table_index += 1
    if not has_text:
        result.warnings.append(
            "PDF без текстового слоя (вероятно, скан). Распознавание невозможно — OCR вне MVP."
        )
    elif not result.tables:
        result.warnings.append("В PDF есть текст, но таблицы не распознаны.")
    return result


# --------------------------------------------------------------------------- #
# Диспетчер по расширению
# --------------------------------------------------------------------------- #
_PARSERS = {
    "xlsx": parse_xlsx,
    "xls": parse_xls,
    "csv": parse_csv,
    "docx": parse_docx,
    "pdf": parse_pdf,
}


def parse(content: bytes, ext: str) -> ParseResult:
    ext = ext.lower().lstrip(".")
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFormat(f"Формат .{ext} не поддерживается")
    return parser(content)
